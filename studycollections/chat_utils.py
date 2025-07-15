import os
from chromadb import PersistentClient

from django.http import HttpResponseForbidden
from django.shortcuts import get_object_or_404
import ollama
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from django.contrib.auth.models import User
from django.core.exceptions import PermissionDenied
from .views.utils import *
from django.contrib import messages
# Initialize ChromaDB client
chroma = PersistentClient(path="./chroma_db")


def chunk_text(text, chunk_size=200, overlap=30):
    """
    Splits text into overlapping chunks.
    """
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunks.append(" ".join(words[i:i + chunk_size]))
    return chunks

def get_embeddings(chunks, model="nomic-embed-text", batch_size=50):
    """
    Generate embeddings for a list of text chunks using Ollama.
    """
    if isinstance(chunks, str):
        chunks = [chunks]

    embeddings = []
    for start in range(0, len(chunks), batch_size):
        batch = chunks[start:start + batch_size]
        try:
            resp = ollama.embed(model=model, input=batch)
            batch_embeds = resp.get("embeddings", [])
        except Exception as e:
            print(f"Error embedding batch at {start}: {e}")
            batch_embeds = [None] * len(batch)
        embeddings.extend(batch_embeds)
    return embeddings

def ingest_document_chunks(doc):
    """
    Reads and chunks an approved document, then embeds and stores in Chroma.
    """
    path = doc.file.path
    ext = os.path.splitext(path)[1].lower()

    # Extract text
    if ext == ".docx":
        docx = DocxDocument(path)
        text = "\n".join([para.text for para in docx.paragraphs])
    elif ext == ".pdf":
        reader = PdfReader(path)
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    else:
        with open(path, encoding="utf-8") as f:
            text = f.read()

    chunks = chunk_text(text)
    embeddings = get_embeddings(chunks)

    # Get the collection
    collection_name = f"collection_{doc.collection.id}"
    db_collection = chroma.get_or_create_collection(name=collection_name)

    for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
        if embedding is None:
            continue  # Skip any invalid embeddings
        metadata = {
            "type": "document",                 # <-- Important: explicitly tag type
            "doc_id": str(doc.id),              # <-- Ensure doc_id is a string
            "chunk_index": idx,
            "title": doc.title,
            "collection_id": str(doc.collection.id),
        }
        db_collection.upsert(
            ids=[f"{doc.id}_{idx}"],
            embeddings=[embedding],
            metadatas=[metadata],
            documents=[chunk],
        )

def delete_document_chunks(user, document_id: int, collection_id: int):
    """
    Deletes all vector chunks from ChromaDB associated with a given document,
    after verifying the user has edit access and the document belongs to the collection.
    """
    from .models import Collection, Document

    collection = get_object_or_404(Collection, id=collection_id)
    if not user_can_edit(user, collection):
        raise PermissionDenied("You do not have permission to modify this collection.")

    document = get_object_or_404(Document, id=document_id)
    if document.collection.id != collection.id:
        raise PermissionDenied("This document does not belong to the specified collection.")

    collection_name = f"collection_{collection_id}"
    db_collection = chroma.get_or_create_collection(name=collection_name)

    # Fetch all chunks from the collection
    results = db_collection.get(include=["ids", "metadatas"])
    chunk_ids = results.get("ids", [])
    metadatas = results.get("metadatas", [])

    print(f"[DEBUG] ChromaDB: {len(chunk_ids)} total chunks found in {collection_name}")

    # Identify chunks to delete
    ids_to_delete = [
        chunk_id for chunk_id, metadata in zip(chunk_ids, metadatas)
        if metadata.get("doc_id") == str(document.id)
    ]

    if ids_to_delete:
        print(f"[DEBUG] Deleting {len(ids_to_delete)} chunks for document ID {document.id}")
        db_collection.delete(ids=ids_to_delete)
    else:
        print(f"[DEBUG] No matching chunks found for document ID {document.id}")

# # with distant
# def get_relevant_context(query: str, user, collection_id: int, top_k: int = 5, min_score: float = 0.2) -> str:
#     """
#     Retrieves top_k most relevant chunks from ChromaDB after verifying user access to the collection.
#     """
#     from studycollections.models import Collection  # Moved inside

#     collection = get_object_or_404(Collection, id=collection_id)
#     if not user_can_view(user, collection):
#         raise PermissionDenied("You do not have access to this collection.")

#     query_embedding = get_embeddings([query])[0]
#     if query_embedding is None:
#         return ""

#     collection_name = f"collection_{collection_id}"
#     db_collection = chroma.get_or_create_collection(name=collection_name)

#     results = db_collection.query(
#         query_embeddings=[query_embedding],
#         n_results=top_k,
#     )

#     documents = results.get("documents", [[]])[0]
#     distances = results.get("distances", [[]])[0]  # assuming this returns cosine distances

#     # Filter based on similarity (convert cosine distance to similarity: 1 - distance)
#     relevant_docs = [
#         doc for doc, dist in zip(documents, distances) if (1 - dist) >= min_score
#     ]

#     return "\n\n".join(relevant_docs)

# Without distant
def get_relevant_context(query: str, user, collection_id: int, top_k: int = 5) -> str:
    """
    Retrieves top_k most relevant chunks from ChromaDB after verifying user access to the collection.
    """
    from studycollections.models import Collection  # Moved inside

    collection = get_object_or_404(Collection, id=collection_id)
    if not user_can_view(user, collection):
        raise PermissionDenied("You do not have access to this collection.")

    query_embedding = get_embeddings([query])[0]
    if query_embedding is None:
        return ""

    collection_name = f"collection_{collection_id}"
    db_collection = chroma.get_or_create_collection(name=collection_name)

    results = db_collection.query(
        query_embeddings=[query_embedding],
        n_results=top_k,
    )

    documents = results.get("documents", [[]])[0]

    return "\n\n".join(documents)

def ingest_public_document_chunks(doc):
    """
    Reads and chunks a public document, then embeds and stores in Chroma.
    """
    path = doc.file.path
    ext = os.path.splitext(path)[1].lower()

    # Extract text
    if ext == ".docx":
        docx = DocxDocument(path)
        text = "\n".join([para.text for para in docx.paragraphs])
    elif ext == ".pdf":
        reader = PdfReader(path)
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    else:
        with open(path, encoding="utf-8") as f:
            text = f.read()

    chunks = chunk_text(text)
    embeddings = get_embeddings(chunks)

    # Save to Chroma using a dedicated public collection
    collection = chroma.get_or_create_collection(name="public_documents")

    for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
        if embedding is None:
            continue
        metadata = {
            "doc_id": str(doc.id),
            "chunk_index": idx,
            "title": doc.title,
            "is_public": True,
        }
        collection.upsert(
            ids=[f"public_{doc.id}_{idx}"],
            embeddings=[embedding],
            metadatas=[metadata],
            documents=[chunk],
        )

def delete_public_document_chunks(doc):
    """
    Removes the chunks related to a public document from ChromaDB.
    """
    collection = chroma.get_or_create_collection(name="public_documents")

    # Query ChromaDB to find all chunk IDs related to this document
    results = collection.get(include=["metadatas"])
    chunk_ids = [
        result_id
        for result_id, metadata in zip(results["ids"], results["metadatas"])
        if metadata.get("doc_id") == str(doc.id)
    ]
    collection.delete(ids=chunk_ids)
    
def get_relevant_public_context(query: str, top_k=10, candidate_k=20, min_score=0.4):
    emb = get_embeddings([query])[0]
    if emb is None: return "", False

    results = chroma.get_or_create_collection("public_documents").query(
        query_embeddings=[emb],
        n_results=candidate_k,
        include=["documents", "distances"]
    )
    docs, dists = results["documents"][0], results["distances"][0]

    # Filter only if similarity >= threshold
    filtered = [
        doc for doc, dist in zip(docs, dists) if dist <= (1 - min_score)
    ]


    chosen = filtered[:top_k] if filtered else docs[:top_k]
    return "\n\n".join(chosen).strip(), bool(filtered)




def ingest_flashcard_to_chromadb(flashcard):
    collection_name = f"collection_{flashcard.collection.id}"
    db_collection = chroma.get_or_create_collection(name=collection_name)

    chunk_id = f"flashcard_{flashcard.id}"
    db_collection.delete(ids=[chunk_id])  # remove old chunk if it exists

    content = f"Flashcard:\nQ: {flashcard.question}\nA: {flashcard.answer}"
    embedding = get_embeddings([content])[0]
    if embedding is not None:
        db_collection.add(
            documents=[content],
            embeddings=[embedding],
            ids=[chunk_id],
            metadatas=[{
                "type": "flashcard",
                "collection_id": flashcard.collection.id,
                "question": flashcard.question,
            }]
        )


def delete_flashcard_from_chromadb(flashcard):
    collection_name = f"collection_{flashcard.collection.id}"
    db_collection = chroma.get_or_create_collection(name=collection_name)
    db_collection.delete(ids=[f"flashcard_{flashcard.id}"])

def ingest_mcq_to_chroma(mcq):
    collection_id = mcq.collection.id
    collection_name = f"collection_{collection_id}"
    db_collection = chroma.get_or_create_collection(name=collection_name)

    choices = [ans.strip() for ans in mcq.answers]
    correct_answers = [choices[i] for i in mcq.correct_indices]
    choices_str = '- Choices: '
    for choice in choices:
        choices_str += f'\n+ {choice}'

    correct_str = '- Correct Answers: '
    for correct_answer in correct_answers:
        correct_str += f'\n+ {correct_answer} '
    
    content = (
        f"Question: {mcq.question_text}\n"
        f"{choices_str}\n"
        f"{correct_str}"
        
    )

    embedding = get_embeddings([content])[0]
    if embedding is None:
        return

    doc_id = f"mcq-{mcq.id}"

    db_collection.upsert(
        documents=[content],
        ids=[doc_id],
        embeddings=[embedding],
        metadatas=[{
            "type": "mcq",
            "collection_id": collection_id,
            "mcq_id": mcq.id
        }],
    )


def delete_mcq_from_chromadb(mcq):
    collection_name = f"collection_{mcq.collection.id}"
    db_collection = chroma.get_or_create_collection(name=collection_name)
    db_collection.delete(ids=[f"mcq_{mcq.id}"])

def ingest_multipart_question_to_chromadb(multipart_question):
    collection = multipart_question.collection
    collection_name = f"collection_{collection.id}"
    db_collection = chroma.get_or_create_collection(name=collection_name)

    parts = []
    for part in multipart_question.parts.all():
        choices = [ans.strip() for ans in part.answers]
        correct_answers = [choices[i] for i in part.correct_indices]
        choices_str = '- Choices: '
        for choice in choices:
            choices_str += f'\n+ {choice}'

        correct_str = '- Correct Answers: '
        for correct_answer in correct_answers:
            correct_str += f'\n+ {correct_answer} '
        
        part_text = (
            f"Question: {part.question_text}\n"
            f"{choices_str}\n"
            f"{correct_str}"
            
        )
        parts.append(part_text)
    document_text = (
        f"Instruction: {multipart_question.instructions}\n\n" + "\n\n".join(parts)
    )
    print(f"DOCUMENT: {document_text}")
    embedding = get_embeddings([document_text])[0]
    print(f"EMBEDDING: {embedding}")

    if embedding is None:
        return

    doc_id = f"multipart_{multipart_question.id}"
    
    db_collection.upsert(
        documents=[document_text],
        ids=[doc_id],
        embeddings=[embedding],
        metadatas=[{
            "type": "multipart_question",
            "collection_id": collection.id,
            "question_id": multipart_question.id,
            "created_by": multipart_question.created_by.id,
        }]
    )



def delete_multipart_question_from_chromadb(multipart_question):
    collection_name = f"collection_{multipart_question.collection.id}"
    db_collection = chroma.get_or_create_collection(name=collection_name)
    db_collection.delete(ids=[f"multipart_{multipart_question.id}"])


def delete_studycollection_from_chromadb(collection_id):
    collection_name = f"collection_{collection_id}"
    chroma.delete_collection(name=collection_name)